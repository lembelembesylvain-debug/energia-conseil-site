#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Générateur de devis ENERGIA-CONSEIL IA® - Rénovation énergétique
Format Word .docx - Conforme MaPrimeRénov' 2026
"""

import sys
import io
from datetime import datetime

if sys.platform == "win32":
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# =============================================================================
# DONNÉES ENTREPRISE
# =============================================================================

ENTREPRISE = {
    "nom": "ENERGIA-CONSEIL IA®",
    "slogan": "Audit Énergétique Intelligence Artificielle",
    "adresse": "16 Rue Cuvier, 69006 Lyon",
    "telephone": "06 10 59 68 98",
    "email": "contact@energia-conseil.com",
    "siret": "94181942700019",
    "rcs": "Lyon 941819427",
    "tva": "FR01941819427",
    "ape": "7112B",
    "capital": "100€",
    "certifications": "FEEBAT | QUALIPV | Auditeur Énergétique",
}

# =============================================================================
# DONNÉES CLIENT EXEMPLE - Mme FONTANAY (modifiable)
# =============================================================================

CLIENT_EXEMPLE = {
    "nom_complet": "Mme FONTANAY Stéphanie",
    "adresse_bien": "167 Rue des Acacias, 42210 MONTROND-LES-BAINS",
    "adresse_fiscale": "",  # Si différent du bien
    "email": "fontanay-stephanie@orange.fr",
    "telephone": "06 59 38 97 49",
    "profil": "BLEU",
    "categorie": "Très modestes",
    "rfr": 24000,
    "parts": 2,
    "pct_aides": 80,
    "commercial": "",
    "contexte": "Maison individuelle, DPE actuel F, surface 103,8 m² ITE + 46 m² plancher. Année construction : années 80. Projet de rénovation globale pour atteindre DPE C.",
    "audit_ref": "",  # Ex: "MUTTA - 18/06/2024 - Valable jusqu'au 06/08/2030"
    "dpe_actuel": "F",
    "dpe_vise": "C",
    "gain_classes": 3,
    "surface_ite": 103.8,
    "surface_plancher": 46,
    "economies_an": 2100,
    "economies_mois": 175,
    "facture_avant": 2800,
    "facture_apres": 700,
    "conso_avant": 280,
    "conso_apres": 110,
    "plus_value": 40000,
    "roi_ans": 12,
    "confort": "Confort thermique amélioré, suppression des ponts thermiques, ventilation maîtrisée.",
    "dpe_expire": False,
    "aides_custom": {"mapr": 31900, "cee": 6500, "region": 1500, "total": 39900},  # Optionnel : montants exacts
    "semaines_total": 12,
    "postes": [
        {
            "numero": 1,
            "nom": "ITE 103,8 m² - Laine de roche",
            "montant_ttc": 18500,
            "description": "Isolation Thermique par l'Extérieur 103,8 m² en laine de roche 140 mm. Enduit minéral rapporté. Bardage ou enduit selon finition. Échafaudage inclus.",
            "artisan": "HB FACADIER",
            "certification": "RGE Qualité Isolation E-E 109538",
            "contact": "Pauline DURY - 07 68 05 74 26",
            "adresse_artisan": "22 Avenue Benoît Fourneyron, 42160 ANDRÉZIEUX-BOUTHÉON",
        },
        {
            "numero": 2,
            "nom": "Isolation plancher bas 46 m²",
            "montant_ttc": 4200,
            "description": "Isolation plancher bas 46 m², sous-face ou sur dalle. Polystyrène ou laine minérale selon configuration.",
            "artisan": "ECO SYSTÈME DURABLE",
            "certification": "RGE Qualité Isolation",
            "contact": "01 70 93 97 15",
            "adresse_artisan": "",
        },
        {
            "numero": 3,
            "nom": "Menuiseries double vitrage",
            "montant_ttc": 9500,
            "description": "Remplacement menuiseries double vitrage argon, 6 fenêtres + 1 porte-fenêtre. PVC ou bois selon modèle choisi.",
            "artisan": "HB FACADIER / C2L",
            "certification": "RGE Menuiseries E-E102269",
            "contact": "",
            "adresse_artisan": "",
        },
        {
            "numero": 4,
            "nom": "PAC Air/Eau 6 kW",
            "montant_ttc": 14000,
            "description": "Pompe à chaleur Air/Eau 6 kW, plancher chauffant ou radiateurs basse température. Raccordement ECS inclus.",
            "artisan": "ECO SYSTÈME DURABLE",
            "certification": "RGE QualiPAC",
            "contact": "01 70 93 97 15",
            "adresse_artisan": "",
        },
        {
            "numero": 5,
            "nom": "Ballon thermodynamique 200 L",
            "montant_ttc": 3500,
            "description": "Ballon thermodynamique 200 L pour eau chaude sanitaire. Raccordement PAC.",
            "artisan": "ECO SYSTÈME DURABLE",
            "certification": "RGE QualiPAC",
            "contact": "01 70 93 97 15",
            "adresse_artisan": "",
        },
        {
            "numero": 6,
            "nom": "Coordination & MAR",
            "montant_ttc": 14363,
            "description": "Coordination projet, constitution dossier MaPrimeRénov' Parcours Accompagné, suivi travaux, réception.",
            "artisan": "ENERGIA-CONSEIL IA®",
            "certification": "AMO Rénovation",
            "contact": "06 10 59 68 98",
            "adresse_artisan": "16 Rue Cuvier, 69006 Lyon",
        },
    ],
    "phase_1": "Administrative – Constitution dossier MPR, validation aides, souscription Éco-PTZ",
    "phase_2": "ITE + menuiseries (6 sem.)",
    "phase_3": "PAC + ballon + plancher (4 sem.)",
    "phase_finale": "Réception finale + PV + DPE post-travaux",
}

# =============================================================================
# RÈGLES MAPRIMERENOV' 2026
# =============================================================================

TAUX_MAPR = {"BLEU": 80, "JAUNE": 60, "VIOLET": 45, "ROSE": 30}
ECRETEMENT_MAX = {"BLEU": 100, "JAUNE": 90, "VIOLET": 80, "ROSE": 50}
PLAFOND_MAPR = {"BLEU": 70000, "JAUNE": 52500, "VIOLET": 42000, "ROSE": 21000}
TVA_REDUITE = 0.055
TVA_10 = 0.10
TVA_20 = 0.20


def set_cell_shading(cell, color):
    """Applique une couleur de fond à une cellule de tableau."""
    shading = parse_xml(
        r'<w:shd {} w:fill="{}"/>'.format(nsdecls("w"), color)
    )
    cell._tc.get_or_add_tcPr().append(shading)


def add_heading_sep(doc, text, level=1):
    """Ajoute un titre avec séparateur."""
    p = doc.add_paragraph()
    p.add_run("═" * 70).bold = False
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = "Arial"
    p = doc.add_paragraph()
    p.add_run("═" * 70).bold = False


def add_paragraph_format(doc, text, bold=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    """Ajoute un paragraphe formaté."""
    p = doc.add_paragraph(style="Normal")
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    run.font.name = "Arial"
    return p


def format_euro(val):
    """Formate un montant en euros."""
    return f"{int(val):,} €".replace(",", " ")


def calculer_aides(client, total_ht):
    """Calcule les aides selon profil 2026. Utilise client['aides_custom'] si fourni."""
    if client.get("aides_custom"):
        a = client["aides_custom"]
        return {
            "mapr": a.get("mapr", 0),
            "cee": a.get("cee", 0),
            "region": a.get("region", 0),
            "total": a.get("total", sum(a.values())),
        }
    profil = client["profil"]
    taux = TAUX_MAPR.get(profil, 60)
    plafond = PLAFOND_MAPR.get(profil, 52500)
    mapr = min(total_ht * taux / 100, plafond)
    cee = 6500 if profil in ("BLEU", "JAUNE") else 3500
    region = 1500
    total = mapr + cee + region
    return {"mapr": mapr, "cee": cee, "region": region, "total": total}


def generer_devis(client_data=None):
    """Génère le document Word du devis."""
    client = client_data or CLIENT_EXEMPLE

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    # ─── HEADER ───
    add_heading_sep(doc, "HEADER (En-tête)", 1)
    p = doc.add_paragraph()
    run = p.add_run(ENTREPRISE["nom"])
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = "Arial"
    add_paragraph_format(doc, ENTREPRISE["slogan"])
    add_paragraph_format(doc, "Marque déposée INPI | Bureau d'études & AMO")
    add_paragraph_format(doc, f"{ENTREPRISE['adresse']} | {ENTREPRISE['telephone']}")
    add_paragraph_format(
        doc,
        f"SIRET : {ENTREPRISE['siret']} | RCS {ENTREPRISE['rcs']} | SASU Capital {ENTREPRISE['capital']}",
    )

    ref_devis = f"I-26-03-{client['nom_complet'].split()[-1].upper()}-RENOV"
    date_jour = datetime.now().strftime("%d/%m/%Y")
    add_paragraph_format(doc, f"DEVIS N° {ref_devis}", bold=True)
    add_paragraph_format(doc, f"Date : {date_jour} | Validité : 30 jours")
    doc.add_paragraph()

    # ─── I. INFORMATIONS CLIENT ───
    add_heading_sep(doc, "I. INFORMATIONS CLIENT")
    add_paragraph_format(doc, f"CLIENT(S) : {client['nom_complet']}", bold=True)
    add_paragraph_format(doc, f"Bien : {client['adresse_bien']}")
    if client.get("adresse_fiscale"):
        add_paragraph_format(doc, f"[Si différent] Adresse fiscale : {client['adresse_fiscale']}")
    add_paragraph_format(doc, f"Contact : {client['email']} | {client['telephone']}")
    add_paragraph_format(
        doc,
        f"Profil MaPrimeRénov' : {client['profil']} ({client['categorie']} – RFR {client['rfr']:,} €, {client['parts']} parts) | {client['pct_aides']}% aides".replace(",", " "),
    )
    if client.get("commercial"):
        add_paragraph_format(doc, f"Interlocuteur commercial : {client['commercial']}")
    doc.add_paragraph()

    # ─── II. CONTEXTE ───
    add_heading_sep(doc, "II. CONTEXTE")
    add_paragraph_format(doc, client["contexte"])
    if client.get("audit_ref"):
        add_paragraph_format(doc, f"Référence audit : {client['audit_ref']}")
    doc.add_paragraph()

    # ─── III. DÉTAIL DES TRAVAUX ───
    postes = client["postes"]
    if client.get("dpe_expire"):
        postes = [
            {
                "numero": 0,
                "nom": "Nouveau DPE réglementaire (si nécessaire)",
                "montant_ttc": 350,
                "description": "Diagnostic de Performance Énergétique réglementaire obligatoire pour constitution dossier MaPrimeRénov'. Bureau d'études certifié. Validité 10 ans.",
                "artisan": "Bureau d'études certifié",
                "certification": "",
                "contact": "",
                "adresse_artisan": "",
            }
        ] + [{**p, "numero": p["numero"] + 1} for p in postes]

    add_heading_sep(doc, f"III. DÉTAIL DES TRAVAUX – {len(postes)} postes")
    if client.get("dpe_expire"):
        add_paragraph_format(doc, "IMPORTANT : Poste 0 inclus car DPE expiré ou absent", bold=True)

    for p in postes:
        add_paragraph_format(doc, f"{p['numero']}. {p['nom']}", bold=True)
        add_paragraph_format(doc, f"{format_euro(p['montant_ttc'])} TTC")
        add_paragraph_format(doc, p["description"])
        if p.get("artisan"):
            add_paragraph_format(
                doc,
                f"✅ ARTISAN RGE – {p['artisan']} | {p.get('certification', '')} | {p.get('contact', '')} | {p.get('adresse_artisan', '')}",
            )
        doc.add_paragraph()

    # ─── IV. TOTAUX ───
    total_ttc = sum(p["montant_ttc"] for p in postes)
    total_ht = round(total_ttc / (1 + TVA_REDUITE), 2)
    tva_55 = round(total_ht * TVA_REDUITE, 2)

    add_heading_sep(doc, "IV. TOTAUX")
    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    rows_totaux = [
        ("Total HT", format_euro(total_ht)),
        ("TVA 5,5 % (isolation/chauffage)", format_euro(tva_55)),
        ("TOTAL TTC", format_euro(total_ttc)),
    ]
    for i, (lib, val) in enumerate(rows_totaux):
        table.rows[i].cells[0].text = lib
        table.rows[i].cells[1].text = val
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = (i == 2)
        table.rows[i].cells[1].paragraphs[0].runs[0].bold = (i == 2)
    doc.add_paragraph()

    # ─── V. AIDES FINANCIÈRES ───
    add_heading_sep(doc, "V. AIDES FINANCIÈRES 2026")
    aides = calculer_aides(client, total_ht)
    add_paragraph_format(
        doc,
        f"Profil {client['profil']} ({client['categorie']}) – RFR {client['rfr']:,} € – {client['parts']} parts".replace(",", " "),
    )

    table_aides = doc.add_table(rows=5, cols=2)
    table_aides.style = "Table Grid"
    rows_aides = [
        (f"MaPrimeRénov' Parcours Accompagné ({client['pct_aides']}%)", format_euro(aides["mapr"])),
        ("CEE " + ("Précarité " if client["profil"] in ("BLEU", "JAUNE") else ""), format_euro(aides["cee"])),
        ("Aide Région Auvergne-Rhône-Alpes", format_euro(aides["region"])),
        ("Total aides", format_euro(aides["total"])),
    ]
    for i, (lib, val) in enumerate(rows_aides):
        table_aides.rows[i].cells[0].text = lib
        table_aides.rows[i].cells[1].text = val
    doc.add_paragraph()

    add_paragraph_format(doc, "━" * 60)
    add_paragraph_format(doc, "⚠️ MENTION IMPORTANTE - VALIDATION DES AIDES", bold=True)
    add_paragraph_format(
        doc,
        "Les montants d'aides indiqués (MaPrimeRénov', CEE, Région) sont des ESTIMATIONS calculées selon les barèmes 2026 et donnés À TITRE INDICATIF.",
    )
    add_paragraph_format(doc, "Ils devront être validés par :")
    add_paragraph_format(doc, "• Le MAR (Mon Accompagnateur Rénov') lors de la constitution du dossier")
    add_paragraph_format(doc, "• L'ANAH lors de l'instruction du dossier MaPrimeRénov' Parcours Accompagné")
    add_paragraph_format(doc, "• Les organismes CEE pour les Certificats d'Économies d'Énergie")
    add_paragraph_format(
        doc,
        "Le montant final peut varier selon l'éligibilité réelle du projet, les plafonds de dépenses applicables et les règles d'écrêtement.",
    )
    add_paragraph_format(
        doc,
        "ENERGIA-CONSEIL s'engage à maximiser les aides dans le cadre réglementaire, mais ne peut garantir les montants exacts avant validation officielle ANAH.",
    )
    add_paragraph_format(doc, "En cas de refus ou de montant inférieur, le client sera informé avant le démarrage des travaux.")
    add_paragraph_format(doc, "━" * 60)
    doc.add_paragraph()

    reste_charge = total_ttc - aides["total"]
    taux_prise = round(aides["total"] / total_ttc * 100) if total_ttc > 0 else 0

    table_recap = doc.add_table(rows=4, cols=2)
    table_recap.style = "Table Grid"
    recap_rows = [
        ("Coût total TTC", format_euro(total_ttc)),
        ("Aides totales estimées", f"-{format_euro(aides['total'])}"),
        ("RESTE À CHARGE ESTIMÉ", format_euro(reste_charge)),
    ]
    for i, (lib, val) in enumerate(recap_rows):
        table_recap.rows[i].cells[0].text = lib
        table_recap.rows[i].cells[1].text = val
        if "RESTE" in lib:
            table_recap.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            table_recap.rows[i].cells[1].paragraphs[0].runs[0].bold = True
    add_paragraph_format(doc, f"Taux de prise en charge estimé : {taux_prise}%")
    doc.add_paragraph()

    # ─── VI. FINANCEMENT Éco-PTZ ───
    add_heading_sep(doc, "VI. FINANCEMENT – Éco-PTZ (0%)")
    mensualite = round(reste_charge / 180, 0) if reste_charge > 0 else 0
    gain_net = client["economies_mois"] - mensualite

    add_paragraph_format(doc, "Éco-PTZ (Éco-Prêt à Taux Zéro) :")
    add_paragraph_format(doc, f"Montant : {format_euro(reste_charge)}")
    add_paragraph_format(doc, "Taux : 0% (aucun intérêt)")
    add_paragraph_format(doc, "Durée : 15 ans (180 mois)")
    add_paragraph_format(doc, f"Mensualité : {format_euro(mensualite)}/mois")
    add_paragraph_format(doc, "Coût total crédit : 0 €")
    doc.add_paragraph()
    add_paragraph_format(doc, "Alternative si Éco-PTZ refusé : VIVONS COURTIER – FABIEN – 06 71 19 96 45")
    add_paragraph_format(doc, "(0 € apport, acceptation 48-72h)")
    doc.add_paragraph()
    add_paragraph_format(
        doc,
        f"Économies vs mensualité : Économies énergie {format_euro(client['economies_mois'])}/mois ({format_euro(client['economies_an'])}/an) − Mensualité {format_euro(mensualite)}/mois = Gain net {'+' if gain_net >= 0 else ''}{format_euro(gain_net)}/mois",
    )
    if gain_net < 0:
        add_paragraph_format(
            doc,
            f"Bilan : Le projet présente un déficit mensuel temporaire de {format_euro(abs(gain_net))}/mois, largement compensé par la plus-value immobilière de +{format_euro(client['plus_value'])} et l'atteinte du DPE {client['dpe_vise']}.",
        )
    doc.add_paragraph()

    # ─── VII. RÉSULTATS ATTENDUS ───
    add_heading_sep(doc, "VII. RÉSULTATS ATTENDUS")
    pct_conso = round((client["conso_avant"] - client["conso_apres"]) / client["conso_avant"] * 100) if client["conso_avant"] else 0
    add_paragraph_format(doc, f"DPE : {client['dpe_actuel']} → {client['dpe_vise']} (+{client['gain_classes']} classes) ✅")
    add_paragraph_format(doc, f"Consommation : {client['conso_avant']} → {client['conso_apres']} kWh/m²/an (-{pct_conso}%)")
    add_paragraph_format(doc, f"Facture : {format_euro(client['facture_avant'])} → {format_euro(client['facture_apres'])}/an (-{format_euro(client['facture_avant'] - client['facture_apres'])}/an)")
    add_paragraph_format(doc, f"Économies : {format_euro(client['economies_an'])}/an ({format_euro(client['economies_mois'])}/mois)")
    add_paragraph_format(doc, f"Gain net mensuel : {'+' if gain_net >= 0 else ''}{format_euro(gain_net)}/mois")
    add_paragraph_format(doc, f"Confort : {client['confort']}")
    add_paragraph_format(doc, f"Plus-value immobilière : +{format_euro(client['plus_value'])}")
    add_paragraph_format(doc, f"ROI : {client['roi_ans']} ans")
    doc.add_paragraph()

    # ─── VIII. PLANNING ───
    add_heading_sep(doc, f"VIII. PLANNING ({client['semaines_total']} semaines)")
    add_paragraph_format(doc, f"Phase 1 (2 sem.) : {client['phase_1']}")
    add_paragraph_format(doc, f"Phase 2 (6 sem.) : {client['phase_2']}")
    add_paragraph_format(doc, f"Phase 3 (4 sem.) : {client['phase_3']}")
    add_paragraph_format(doc, "Phase finale (1 j) : Réception finale + PV + DPE post-travaux")
    doc.add_paragraph()

    # ─── IX. MENTIONS LÉGALES ───
    add_heading_sep(doc, "IX. MENTIONS LÉGALES")
    add_paragraph_format(
        doc,
        f"Devis établi par {ENTREPRISE['nom']} – SASU au capital de {ENTREPRISE['capital']} – SIRET {ENTREPRISE['siret']} – RCS {ENTREPRISE['rcs']} – N° TVA {ENTREPRISE['tva']} – APE {ENTREPRISE['ape']}. Marque déposée INPI.",
    )
    add_paragraph_format(doc, "Paiement : 30% à la commande, 40% à mi-parcours, 30% à réception.")
    add_paragraph_format(doc, "RGPD : Données traitées pour la gestion du projet ; contact contact@energia-conseil.com.")
    add_paragraph_format(doc, "En cas de litige : Médiation MEDICYS - www.medicys.fr")
    doc.add_paragraph()

    # ─── X. BON POUR ACCORD ───
    add_heading_sep(doc, "X. BON POUR ACCORD")
    add_paragraph_format(
        doc,
        f"Je soussigné(e) {client['nom_complet']}, accepte le présent devis pour un montant de {format_euro(total_ttc)} TTC et mandate {ENTREPRISE['nom']} pour la coordination du projet de rénovation énergétique.",
    )
    add_paragraph_format(doc, "Fait à __________________________ , le __________________________")
    add_paragraph_format(doc, "Signature : _______________________________________")
    doc.add_paragraph()

    # ─── FOOTER ───
    add_heading_sep(doc, "FOOTER")
    add_paragraph_format(doc, f"{ENTREPRISE['nom']} - Marque déposée INPI")
    add_paragraph_format(doc, f"SIRET : {ENTREPRISE['siret']} | RCS {ENTREPRISE['rcs']} | N° TVA : {ENTREPRISE['tva']}")
    add_paragraph_format(doc, f"Code APE : {ENTREPRISE['ape']} | SASU Capital {ENTREPRISE['capital']}")
    add_paragraph_format(doc, "Paiement : 30% / 40% / 30% | RGPD : contact@energia-conseil.com")
    add_paragraph_format(doc, "Médiation : MEDICYS - www.medicys.fr")
    add_paragraph_format(doc, "© 2026 ENERGIA-CONSEIL IA - Tous droits réservés")

    return doc, total_ttc


def main():
    print("Génération du devis ENERGIA-CONSEIL IA®...")
    doc, total = generer_devis()
    filename = f"Devis_ENERGIA_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    doc.save(filename)
    print(f"✅ Devis créé : {filename}")
    print(f"   Montant total TTC : {total:,} €".replace(",", " "))


if __name__ == "__main__":
    main()
